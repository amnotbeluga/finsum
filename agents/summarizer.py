import os
import PyPDF2
import docx
import pandas as pd
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import warnings
import re
warnings.filterwarnings('ignore')

class FinancialDocumentSummarizer:
    def __init__(self, model_name="google/pegasus-xsum", skip_pages=2):
        print(f"Loading model: {model_name}")
        self.model_name = model_name
        self.skip_pages = skip_pages
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        self.summarizer = pipeline("summarization", model=self.model, tokenizer=self.tokenizer)
    
    def is_table_content(self, text, threshold=0.3):
        if not text or len(text.strip()) < 50:
            return False
        
        table_indicators = [
            r'\b\d+\.\d+\b',
            r'\$\s*\d+[,.]?\d*',
            r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b',
            r'\|',
            r'\t',
            r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b',
            r'\bQ[1-4]\s*\d{4}\b',
            r'\bFY\s*\d{4}\b',
        ]
        multiple_spaces = len(re.findall(r' {2,}', text))
        pattern_matches = 0
        for pattern in table_indicators:
            matches = len(re.findall(pattern, text))
            if matches > 2:
                pattern_matches += matches
        lines = text.split('\n')
        if lines:
            avg_line_length = sum(len(line) for line in lines) / len(lines)
            consistent_line_lengths = all(abs(len(line) - avg_line_length) < 20 for line in lines[:5])
        else:
            consistent_line_lengths = False
        if pattern_matches > 10 or multiple_spaces > 20:
            return True
        if consistent_line_lengths and pattern_matches > 5:
            return True
        
        return False
    
    def clean_text(self, text):
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\$\.,%\(\)\/\-]', '', text)
        text = re.sub(r'\s+([.,;:)])', r'\1', text)
        text = re.sub(r'([(])\s+', r'\1', text)
        
        return text.strip()
    
    def extract_text_from_pdf(self, pdf_path):
        all_text = []
        skipped_pages = []
        table_pages = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                print(f"Total pages in PDF: {total_pages}")
                print(f"Skipping first {self.skip_pages} pages...")
                
                for page_num in range(total_pages):
                    if page_num < self.skip_pages:
                        skipped_pages.append(page_num + 1)
                        continue
                    
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if not page_text or len(page_text.strip()) < 50:
                        print(f"  Page {page_num + 1}: Empty or too short, skipping")
                        continue
                    
                    if self.is_table_content(page_text):
                        print(f"  Page {page_num + 1}: Table detected, skipping")
                        table_pages.append(page_num + 1)
                        continue
                    
                    cleaned_text = self.clean_text(page_text)
                    
                    if len(cleaned_text.split()) > 50:
                        all_text.append(cleaned_text)
                        print(f"  Page {page_num + 1}: Added ({len(cleaned_text.split())} words)")
                    else:
                        print(f"  Page {page_num + 1}: Too few words, skipping")
                
                print(f"\nProcessing summary:")
                print(f"  - Skipped pages (first {self.skip_pages}): {skipped_pages}")
                print(f"  - Table pages skipped: {table_pages}")
                print(f"  - Pages processed: {len(all_text)}")
                
        except Exception as e:
            print(f"Error reading PDF: {e}")
        
        return ' '.join(all_text)
    
    def extract_text_from_docx(self, docx_path):
        all_text = []
        table_count = 0
        
        try:
            doc = docx.Document(docx_path)
            
            start_para = min(self.skip_pages * 3, len(doc.paragraphs))
            
            for i, paragraph in enumerate(doc.paragraphs):
                if i < start_para:
                    continue
                
                para_text = paragraph.text.strip()
                
                if para_text and len(para_text.split()) > 10:

                    if self.is_table_content(para_text):
                        table_count += 1
                        continue
                    
                    cleaned_text = self.clean_text(para_text)
                    all_text.append(cleaned_text)
            
            for table in doc.tables:
                table_count += 1
            
            print(f"DOCX processing: Skipped {table_count} tables/headers")
            
        except Exception as e:
            print(f"Error reading DOCX: {e}")
        
        return ' '.join(all_text)
    
    def read_document(self, file_path):
        file_extension = os.path.splitext(file_path)[1].lower()
        
        print(f"\nReading {file_extension.upper()} file...")
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def extract_text_from_txt(self, txt_path):
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
                lines = text.split('\n')
                if len(lines) > self.skip_pages * 30:
                    text = '\n'.join(lines[self.skip_pages * 30:])
                return self.clean_text(text)
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ""
    
    def safe_summarize(self, text, max_length=None, min_length=None):
        if not text or len(text.split()) < 20:
            return "Insufficient text for summarization"
        
        try:
            text_length = len(text.split())
            
            if max_length is None:
                max_length = min(250, int(text_length * 0.5))
            
            if min_length is None:
                min_length = max(20, int(max_length * 0.3))
            
            if max_length >= text_length:
                max_length = int(text_length * 0.8)
                min_length = int(max_length * 0.3)
            
            max_length = max(max_length, 30)
            min_length = max(min_length, 10)
            
            print(f"  Summarizing: {text_length} words -> {max_length} words")
            
            summary = self.summarizer(
                text, 
                max_length=max_length, 
                min_length=min_length,
                do_sample=False
            )[0]['summary_text']
            
            return summary
            
        except IndexError as e:
            print(f"  IndexError in summarization: {e}")
            sentences = text.split('.')[:3]
            return '. '.join(sentences) + '.'
            
        except Exception as e:
            print(f"  Error in summarization: {e}")
            return "Error generating summary"
    
    def chunk_and_summarize(self, text, target_ratio=0.3):
        words = text.split()
        
        if len(words) < 100:
            return self.safe_summarize(text)
        
        chunk_size = min(500, max(200, len(words) // 10))
        
        chunks = []
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i+chunk_size])
            if len(chunk.split()) > 50:
                chunks.append(chunk)
        
        print(f"Split into {len(chunks)} chunks")
        
        summaries = []
        for i, chunk in enumerate(chunks):
            print(f"  Processing chunk {i+1}/{len(chunks)}")
            summary = self.safe_summarize(chunk, max_length=150, min_length=40)
            if summary and "Insufficient" not in summary and "Error" not in summary:
                summaries.append(summary)
        
        if not summaries:
            return "Could not generate summary from document"
        
        combined = ' '.join(summaries)
        if len(combined.split()) > 300:
            final_summary = self.safe_summarize(combined, max_length=250, min_length=100)
            return final_summary
        
        return combined
    
    def extract_financial_highlights(self, text, max_highlights=5):
        if not text or len(text.split()) < 50:
            return ["Insufficient text for highlights"]
        
        highlights = []
        
        financial_patterns = [
            (r'\b(?:revenue|sales|turnover)\s+(?:was|is|reached|increased|decreased)\s+[\$\€\£]?\s*\d+\.?\d*\s*(?:million|billion|bn|mn)?', 'Revenue'),
            (r'\b(?:net\s+)?(?:profit|loss|income)\s+(?:was|is)\s+[\$\€\£]?\s*\d+\.?\d*\s*(?:million|billion)?', 'Profit/Loss'),
            (r'\bEPS\s+(?:was|is)\s+[\$\€\£]?\s*\d+\.?\d*', 'EPS'),
            (r'\b(?:EBITDA|operating\s+income)\s+(?:was|is)\s+[\$\€\£]?\s*\d+\.?\d*\s*(?:million|billion)?', 'EBITDA'),
            (r'\b(?:dividend|payout)\s+(?:was|is)\s+[\$\€\£]?\s*\d+\.?\d*', 'Dividend'),
            (r'\b(?:cash|reserves)\s+(?:was|is|stood\s+at)\s+[\$\€\£]?\s*\d+\.?\d*\s*(?:million|billion)?', 'Cash'),
            (r'\b(?:Q[1-4]|quarter)\s+\d{4}\s+(?:results|performance)', 'Quarterly Results'),
        ]
        
        sentences = text.split('. ')
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20:
                continue
                
            for pattern, label in financial_patterns:
                if re.search(pattern, sentence, re.IGNORECASE):
                    highlights.append(f"[{label}] {sentence[:150]}...")
                    break
        
        seen = set()
        unique_highlights = []
        for h in highlights:
            if h not in seen and len(unique_highlights) < max_highlights:
                seen.add(h)
                unique_highlights.append(h)
        
        return unique_highlights if unique_highlights else ["No specific financial highlights found"]
    
    def generate_comprehensive_summary(self, file_path):
        try:
            print(f"\n{'='*60}")
            print(f"PROCESSING: {os.path.basename(file_path)}")
            print('='*60)
            
            text = self.read_document(file_path)
            
            if not text or len(text.split()) < 50:
                return {
                    'error': True,
                    'message': 'Document contains insufficient text for summarization',
                    'file_name': os.path.basename(file_path)
                }
            
            word_count = len(text.split())
            print(f"\nTotal words extracted: {word_count}")
            
            print("\nGenerating summary...")
            summary = self.chunk_and_summarize(text)
            
            print("Extracting financial highlights...")
            highlights = self.extract_financial_highlights(text)
            
            report = {
                'error': False,
                'file_name': os.path.basename(file_path),
                'file_size': f"{os.path.getsize(file_path) / 1024:.2f} KB",
                'document_length': f"{word_count} words",
                'summary': summary,
                'summary_word_count': len(summary.split()),
                'financial_highlights': highlights,
                'pages_skipped': self.skip_pages
            }
            
            return report
            
        except Exception as e:
            print(f"Error in generate_comprehensive_summary: {e}")
            return {
                'error': True,
                'message': str(e),
                'file_name': os.path.basename(file_path)
            }

def main():
    summarizer = FinancialDocumentSummarizer(
        model_name="google/pegasus-xsum",
        skip_pages=1
    )
    file_path = input("Enter the path to your financial document: ").strip()
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    report = summarizer.generate_comprehensive_summary(file_path)
    
    if report.get('error'):
        print(f"\nError: {report['message']}")
        return
    
    print("\n" + "="*60)
    print("FINANCIAL DOCUMENT SUMMARY REPORT")
    print("="*60)
    print(f"File: {report['file_name']}")
    print(f"Size: {report['file_size']}")
    print(f"Original length: {report['document_length']}")
    print(f"Summary length: {report['summary_word_count']} words")
    print(f"Pages skipped: {report['pages_skipped']}")
    
    print("\n" + "-"*60)
    print("SUMMARY:")
    print("-"*60)
    print(report['summary'])

    
    save_option = input("\nSave summary to file? (y/n): ").strip().lower()
    if save_option == 'y':
        output_file = f"summary_{os.path.splitext(report['file_name'])[0]}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("="*60 + "\n")
            f.write("FINANCIAL DOCUMENT SUMMARY REPORT\n")
            f.write("="*60 + "\n\n")
            f.write(f"File: {report['file_name']}\n")
            f.write(f"Size: {report['file_size']}\n")
            f.write(f"Original: {report['document_length']}\n")
            f.write(f"Summary: {report['summary_word_count']} words\n\n")
            f.write("-"*60 + "\n")
            f.write("SUMMARY:\n")
            f.write("-"*60 + "\n")
            f.write(report['summary'] + "\n\n")
        
        print(f"\nSummary saved to: {output_file}")

if __name__ == "__main__":
    main()